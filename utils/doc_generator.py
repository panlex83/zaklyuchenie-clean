# utils/doc_generator.py

from docx import Document
from docx.shared import Inches
from datetime import datetime
import os
import json

from utils.gpt_photo_analysis_batch import analyze_photos_batch
from utils.gpt_conclusions import generate_conclusions

# Если "defects" иногда возвращается списком — приводим к строке
def _defects_to_str(d):
    if isinstance(d, list):
        return "; ".join(str(e).strip() for e in d if e)
    return str(d or "")

ELEMENT_LABELS = {
    "facade":    "Фасад",
    "foundation": "Фундамент",
    "walls":      "Стены",
    "roof":       "Кровля",
    "windows":    "Окна и двери",
}
PHOTO_LABELS = {
    "facade":    "Фасад",
    "foundation":"Фундамент",
    "walls":      "Стены",
    "roof":       "Кровля",
    "windows":    "Окна и двери",
}

def generate_doc(data: dict) -> str:
    doc = Document()

    # — Шапка с логотипом
    section = doc.sections[0]
    header = section.header
    logo = os.path.join(os.getcwd(), "logo.png")
    hdr = header.add_paragraph()
    if os.path.exists(logo):
        run = hdr.add_run()
        run.add_picture(logo, width=Inches(1.0))
    text_run = hdr.add_run("\nТЕХНИЧЕСКОЕ ЗАКЛЮЧЕНИЕ № QAZ-__/__-__")
    text_run.bold = True
    hdr.alignment = 1

    # — Реквизиты
    fio     = data.get("full_name", "не распознано")
    id_no   = data.get("id_number", "")
    id_date = data.get("id_date", "")
    addr    = data.get("address", "не распознан")
    cad     = data.get("cadastral_number", "не указан")
    year    = data.get("build_year", "не указан")
    purp    = data.get("purpose", "не указано")
    comp    = data.get("complexity_category", "не определена")

    p = doc.add_paragraph()
    p.add_run("Специализированная организация").bold = True
    p.add_run(" – ТОО «Qazlife», БИН 171140033560,")
    p.add_run(" свидетельство об аккредитации №").bold = True
    p.add_run(" KZ23VWC00067014 от 09 апреля 2024 года, на право осуществления экспертных работ по техническому обследованию надежности и устойчивости зданий и сооружений на технически и технологически сложных объектах первого и второго уровней ответственности,")
    p.add_run(" в лице экспертов").bold = True
    p.add_run(" – Капас Асет Саулебекулы и Жалғасбай Қалдыбек Рысбекұлы, ")
    p = doc.add_paragraph()
    p.add_run("аттестат эксперта").bold = True
    p.add_run(" – Капас Асет Саулебекулы № KZ14VJE00052616; по экспертным работам и инжиниринговым услугам с правом осуществления этой деятельности: по виду: Техническое обследование надежности и устойчивости зданий и сооружений")
    p = doc.add_paragraph()
    p.add_run("дата выдачи").bold = True
    p.add_run(" – 04.02.2020 года.")
    paragraph = doc.add_paragraph()
    paragraph.add_run("аттестат эксперта").bold = True
    paragraph.add_run(" – Жалғасбай Қалдыбек Рысбекұлы № KZ13VJE00050985 по экспертным работам и инжиниринговым услугам с правом осуществления этой деятельности: по виду: Техническое обследование надежности и устойчивости зданий и сооружений")
    paragraph = doc.add_paragraph()
    paragraph.add_run("дата выдачи").bold = True
    paragraph.add_run(" – 25.11.2019 года.")
    paragraph = doc.add_paragraph()
    paragraph.add_run("    Инженер по обследованию").bold = True
    paragraph.add_run(" – Кірбасов Еркебұлан Рүстемұлы.")
    paragraph = doc.add_paragraph()
    paragraph.add_run("Произвели:").bold = True
    paragraph.add_run(f" техническое обследование надежности и устойчивости объекта по адресу: {addr}")

    paragraph = doc.add_paragraph()
    paragraph.add_run("Причина обследования").bold = True
    paragraph.add_run(f" – обращение заявителя {fio}.")
    paragraph = doc.add_paragraph()
    paragraph.add_run("Перечень выполненных работ:").bold = True
    doc.add_paragraph("Обследование объекта с оценкой при необходимости живучести производился специализированной организацией, имеющей аттестованных экспертов по техническому обследованию надежности и устойчивости зданий и сооружений.")
    doc.add_paragraph("    Обследование объекта состоит из этапов:")
    doc.add_paragraph("    - подготовки к проведению обследования;")
    doc.add_paragraph("    - предварительного визуального и полного (детального инструментального) обследования;")
    doc.add_paragraph("    - сплошное визуальное обследование конструкции жилого дома;")
    doc.add_paragraph("    - выявление дефектов и повреждений конструкций с фотофиксацией;")
    doc.add_paragraph("    - определение конструктивного решения;")
    doc.add_paragraph("    - оценки технического состояния и (при необходимости) живучести объекта на аварийное воздействие.")
    doc.add_paragraph("    - составление технического отчета с выводами о техническом состоянии и рекомендациями по дальнейшей эксплуатации.")
    doc.add_paragraph("    На этапе подготовки к проведению обследования выполнены работы по:")
    doc.add_paragraph("    – ознакомлению с объектом обследования, его объемно-планировочным и конструктивным решением;")
    doc.add_paragraph("    На этапе предварительного визуального и полного (детального инструментального) обследования произведена предварительная оценка технического состояния строительных конструкций по внешним признакам для определения необходимости в проведении детального инструментального обследования.")
    doc.add_paragraph("    Предварительную оценку технического состояния строительных конструкций и инженерных систем по внешним признакам следует производить для оперативного выявления явно аварийных участков и своевременного выполнения страховочных мероприятий.")
    doc.add_paragraph("    Если в процессе предварительного обследования будут обнаружены дефекты и повреждения, снижающие прочность, устойчивость и жесткость несущих конструкций, или приводящие к неисправности инженерных систем, то необходимо перейти к детальному инструментальному обследованию.")

    doc.add_heading("    Строительные конструкции:", level=1)
    doc.add_paragraph(f"    Год постройки – {year}")
    doc.add_paragraph(f"    Фундамент – ")
    doc.add_paragraph(f"    Стены – ")
    doc.add_paragraph(f"    Покрытия – ")
    doc.add_paragraph(f"    Покрытие пола – ")
    doc.add_paragraph(f"    Дверные и оконные блоки – ")
    doc.add_paragraph(f"    Покрытие кровли – ")
    doc.add_paragraph(f"    Инженерные сети – ")
    
    paragraph = doc.add_paragraph()
    paragraph.add_run("    Все конструкции находятся ")
    paragraph.add_run("в аварийном неудовлетворительном ").bold = True
    paragraph.add_run("техническом состоянии.")

    doc.add_paragraph(f"Заказчик: {fio}")
    doc.add_paragraph(f"Удостоверение личности: {id_no} от {id_date}")
    doc.add_paragraph(f"Адрес объекта: {addr}")
    doc.add_paragraph(f"Кадастровый номер: {cad}")
    doc.add_paragraph(f"Год постройки: {year}")
    doc.add_paragraph(f"Назначение: {purp}")

    # — Анализ конструктивов
   #doc.add_heading("Описание конструктивных элементов", level=1)
    analysis = {}
    photos_by_type = {}

    for key, label in ELEMENT_LABELS.items():
        photos = data.get(key, []) or []
        photos_by_type[key] = photos

        if not photos:
            doc.add_paragraph(f"{label}: фото не предоставлены")
            continue

        result = analyze_photos_batch(photos, label)

        # fallback: если пакетный анализ вернул [], пробуем каждый фото по отдельности
        if not result and len(photos) > 1:
            result = []
            for p in photos:
                sub = analyze_photos_batch([p], label)
                if sub:
                    result.extend(sub)

        analysis[key] = result

        if result:
            for obj in result:
                idx = obj.get("index", "?")
                desc  = obj.get("description", "").strip()
                defs  = _defects_to_str(obj.get("defects", ""))
                #doc.add_paragraph(f"{label} (фото {idx}): {desc}; Дефекты: {defs or 'не указаны'}")
        else:
            doc.add_paragraph(f"{label}: GPT‑анализ не вернул данных")

    # — Фотофиксация с подписями
    doc.add_heading("Фотофиксация", level=1)

    for key, label in PHOTO_LABELS.items():
        photos = photos_by_type.get(key, [])
        arr    = analysis.get(key, [])

        for idx, path in enumerate(photos):
            if path and os.path.exists(path):
                doc.add_paragraph(f"{label} (фото {idx})")
                doc.add_picture(path, width=Inches(5.5))
                if idx < len(arr):
                    obj  = arr[idx]
                    desc = obj.get("description", "").strip()
                    defs = _defects_to_str(obj.get("defects", ""))
                    overal = obj.get("overall_state", "").strip()
                    doc.add_paragraph(f"Анализ: {desc}; Дефекты: {defs or '—'}; Состояние: {overal or '—'}")

    # — Выводы и рекомендации
    # --- анализ завершён, передаём данные для итогов ---
    data["analysis"] = analysis
    data["element_labels"] = ELEMENT_LABELS  # добавляем словарь для доступа в generate_conclusions

    # --- теперь генерируем раздел "Выводы и рекомендации" ---
    cons = generate_conclusions(data)
    doc.add_heading("Общие выводы:", level=1)

    doc.add_paragraph("    Согласно результатом обследование строительных конструкций, выявлено следущее:")
    paragraph = doc.add_paragraph()
    paragraph.add_run("    - Фундаменты соответствуют категории – ")
    found = analysis.get("foundation", [{}])[0].get("overall_state", "неудовлетворительное")
    paragraph.add_run(f"{found}").bold = True
    paragraph.add_run(" существуют повреждения, свидетельствующие о возможности обрушения конструкции. Требуется немедленная разгрузка конструкций;")
    paragraph = doc.add_paragraph()
    paragraph.add_run("    - Стены соответствуют категории – ")
    wall = analysis.get("walls", [{}])[0].get("overall_state", "неудовлетворительное")
    paragraph.add_run(f"{wall}").bold = True
    paragraph.add_run(" полное повреждение. Снижение несущей способности до 50%. В конструкциях наблюдаются деформации и дефекты, свидетельствующие о потере ими несущей способности. Состояние конструкций аварийное. Возникает угроза обрушения. Необходимо запретить эксплуатацию аварийных конструкций, прекратить технологический процесс и немедленно удалить людей из опасных зон. Конструкция подлежит разборке;")
    paragraph = doc.add_paragraph()
    paragraph.add_run("    - Фасад соответствуют категории – ")
    facade = analysis.get("facade", [{}])[0].get("overall_state", "неудовлетворительное")
    paragraph.add_run(f"{facade}").bold = True
    paragraph.add_run(" существуют повреждения, свидетельствующие о возможности обрушения конструкции. Требуется немедленная разгрузка конструкций;")
    paragraph = doc.add_paragraph()
    paragraph.add_run("    - Дверные и оконные блоки соответствуют категории – ")
    win = analysis.get("windows", [{}])[0].get("overall_state", "неудовлетворительное")
    paragraph.add_run(f"{win}").bold = True
    paragraph.add_run(" техническому состоянию;")
    paragraph = doc.add_paragraph()
    paragraph.add_run("    - Покрытия кровли соответствуют категории – ")
    roof_state = analysis.get("roof", [{}])[0].get("overall_state", "неудовлетворительное")
    paragraph.add_run(f"{roof_state}").bold = True
    paragraph.add_run(" техническому состоянию.")
    paragraph = doc.add_paragraph()
    paragraph.add_run("    Детальное обследование показало, что на объекте ")
    paragraph.add_run("существуют повреждения, свидетельствующие о возможности обрушения конструкции.").bold = True
    paragraph = doc.add_paragraph()
    paragraph.add_run(" Согласно п. 4.3.2, СП РК 1.04-101-2012, объект по вышеуказанному адресу, соответствует — ")
    paragraph.add_run("на грани обрушения,").bold = True
    paragraph.add_run(" характеризуется повреждениями и деформациями, свидетельствующими об исчерпании несущей способности и опасности обрушения (необходимо проведение срочных страховочных мероприятий).")
    paragraph = doc.add_paragraph()
    paragraph.add_run("    В конструкциях наблюдаются деформации и дефекты, свидетельствующие о потере ими несущей способности. Состояние конструкций ")
    paragraph.add_run("аварийное.").bold = True
    paragraph.add_run(" Возникает угроза обрушения. Необходимо запретить эксплуатацию аварийных конструкций, прекратить технологический процесс и немедленно удалить людей из опасных зон. Конструкция подлежит разборке.")
    paragraph = doc.add_paragraph()
    paragraph.add_run(f"    Объект обследования по адресу: {addr} ")
    paragraph.add_run("не может быть допущен к эксплуатации,").bold = True
    paragraph.add_run(" на установленных параметрах. Необходимо ")
    paragraph.add_run("произвести снос здания.").bold = True

    #doc.add_paragraph(f"Общее состояние: {cons['overall_state']}")
    #doc.add_paragraph(f"Выявленные дефекты: {cons['defects']}")
    #doc.add_paragraph(f"Рекомендации: {cons['recommendations']}")

    # — Подписи экспертов
    doc.add_paragraph("   Инженер‑эксперт Капас А.С.")
    doc.add_paragraph("   (Аттестат № KZ14VJE00052616 ")
    paragraph = doc.add_paragraph()
    paragraph.add_run("   от 04.02.2020 г.)         __________________________")
    paragraph.add_run("Капас А.С").bold = True
    doc.add_paragraph("   Инженер‑эксперт Жалғасбай Қ.Р.")
    doc.add_paragraph("   (Аттестат № KZ13VJE00050985 ")
    paragraph = doc.add_paragraph()
    paragraph.add_run("   от 25.11.2019 г.)         __________________________")
    paragraph.add_run("Жалғасбай Қ.Р").bold = True
    paragraph = doc.add_paragraph()
    paragraph.add_run("   Инженер по обследованию         __________________________")
    paragraph.add_run("Кірбасов Е.Р.").bold = True

    os.makedirs("output", exist_ok=True)
    filename = f"output/zaklyuchenie_{datetime.now():%Y%m%d_%H%M%S}.docx"
    doc.save(filename)
    return filename
